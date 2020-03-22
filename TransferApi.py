import docx
from fpdf import FPDF

import Patient as p

class PDF(FPDF):
    def header(self):
        # Arial bold 15
        self.set_font('Arial', 'B', 20)
        # Move to the right
        self.cell(80)
        # Title
        self.cell(30, 10, 'Patient Information', 0, 0, 'C')
        # Line break
        self.ln(20)
        self.set_text_color(220, 50, 50)

    # Page footer
    def footer(self):
        # Position at 1.5 cm from bottom
        self.set_y(-15)
        # Arial italic 8
        self.set_font('Arial', 'I', 8)
        # Page number
        self.cell(0, 10, 'Page ' + str(self.page_no()) + '/{nb}', 0, 0, 'C')


def intergrate_information(patient):
    infor = []
    id = " ".join(["ID:", patient.id])
    infor.append(id)
    gender = " ".join(["Gender:", patient.gender])
    infor.append(gender)
    birth_date = " ".join(["Birth Date:", patient.birth_date])
    infor.append(birth_date)
    marital_status = " ".join(["Marital Status:", patient.marital_status])
    infor.append(marital_status)
    communication = " ".join(["Communication:", patient.communication])
    infor.append(communication)
    if(patient.address is not None):
        lines = " ".join(patient.address.line)
        add = ", ".join([lines, patient.address.city, patient.address.state, patient.address.country, patient.address.postalCode])
        address = " ".join(["Address:", add])
        infor.append(address)
    identifiers = patient.identifier
    for identifier in identifiers:
        identifier_infor = ": ".join([identifier.type, identifier.value])
        infor.append(identifier_infor)
    return infor


def testJSONtoWord(patients):
    doc = docx.Document()
    i = 1
    for patient in patients:
        infor_list = intergrate_information(patient)
        doc.add_heading(" ".join(["Patient", str(i)]), 0)
        for j in range(0, len(infor_list)):
            doc.add_paragraph(infor_list[j])
        i = i+1

    doc.save('Patient.docx')


def testJSONtoPdf(patients):
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    j = 0
    for patient in patients:
        j = j+1
        infor_list = intergrate_information(patient)
        patient_number = " ".join(["Patient", str(j)])
        pdf.set_font('Times', 'B', 15)
        pdf.cell(0, 10, patient_number, 1, 1)
        for i in range(0, len(infor_list)):
            pdf.set_font('Times', '', 12)
            pdf.cell(0, 10, infor_list[i], 0, 1)
        pdf.ln(4)
    pdf.output('Patient.pdf', 'F')

